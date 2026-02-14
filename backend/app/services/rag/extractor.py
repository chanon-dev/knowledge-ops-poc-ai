import csv
import io
from pathlib import Path


class DocumentExtractor:
    """Extract text content from various document formats."""

    SUPPORTED_TYPES = {
        "application/pdf": "_extract_pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "_extract_docx",
        "text/plain": "_extract_text",
        "text/markdown": "_extract_text",
        "text/csv": "_extract_csv",
        "text/html": "_extract_html",
    }

    def extract(self, file_path: str, mime_type: str) -> str:
        method_name = self.SUPPORTED_TYPES.get(mime_type)
        if not method_name:
            ext = Path(file_path).suffix.lower()
            ext_map = {
                ".pdf": "_extract_pdf",
                ".docx": "_extract_docx",
                ".txt": "_extract_text",
                ".md": "_extract_text",
                ".csv": "_extract_csv",
                ".html": "_extract_html",
                ".htm": "_extract_html",
            }
            method_name = ext_map.get(ext)

        if not method_name:
            raise ValueError(f"Unsupported file type: {mime_type} ({file_path})")

        method = getattr(self, method_name)
        return method(file_path)

    def _extract_pdf(self, file_path: str) -> str:
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
            return "\n\n".join(pages)
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF extraction: pip install PyPDF2")

    def _extract_docx(self, file_path: str) -> str:
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            return "\n\n".join(paragraphs)
        except ImportError:
            raise ImportError("python-docx is required for DOCX extraction: pip install python-docx")

    def _extract_text(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _extract_csv(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            rows = []
            for row in reader:
                rows.append(" | ".join(row))
            return "\n".join(rows)

    def _extract_html(self, file_path: str) -> str:
        try:
            from bs4 import BeautifulSoup

            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                soup = BeautifulSoup(f.read(), "html.parser")

            for tag in soup(["script", "style", "nav", "header", "footer"]):
                tag.decompose()

            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                import re
                html = f.read()
                text = re.sub(r"<[^>]+>", " ", html)
                return re.sub(r"\s+", " ", text).strip()
